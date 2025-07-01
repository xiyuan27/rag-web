#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
from functools import reduce
from io import BytesIO
from timeit import default_timer as timer

from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from markdown import markdown 
from PIL import Image
from tika import parser

from api.db import LLMType
from api.db.services.llm_service import LLMBundle
from deepdoc.parser import DocxParser, ExcelParser, HtmlParser, JsonParser, MarkdownParser, PdfParser, TxtParser
from deepdoc.parser.figure_parser import VisionFigureParser, vision_figure_parser_figure_data_wrapper
from deepdoc.parser.pdf_parser import PlainParser, VisionParser
from rag.nlp import concat_img, find_codec, naive_merge, naive_merge_with_images, naive_merge_docx, rag_tokenizer, tokenize_chunks, tokenize_chunks_with_images, tokenize_table


class ImprovedMarkdown(MarkdownParser):
    """改进的Markdown解析器，支持层次化章节切分"""
    
    def __init__(self, chunk_token_num=128):
        super().__init__(chunk_token_num)
        self.heading_stack = []  # 标题栈，维护当前层次
    
    def parse_heading_level(self, line):
        """解析标题级别"""
        line = line.strip()
        if not line.startswith('#'):
            return None, None
        
        # 计算#的数量
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break
        
        if level > 6:  # Markdown最多支持6级标题
            level = 6
            
        # 提取标题文本
        title = line[level:].strip()
        return level, title
    
    def get_chapter_path(self, filename):
        """构建章节路径"""
        # 获取文件名（去除扩展名）
        file_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not file_name:
            file_name = "未命名文档"
        
        # 构建层次路径
        path_parts = [file_name]
        for level, title in self.heading_stack:
            path_parts.append(title)
        
        return " > ".join(path_parts)
    
    def update_heading_stack(self, level, title):
        """更新标题栈，维护层次结构"""
        # 移除比当前级别深的标题
        while self.heading_stack and self.heading_stack[-1][0] >= level:
            self.heading_stack.pop()
        
        # 添加当前标题
        self.heading_stack.append((level, title))
    
    def process_markdown_sections(self, markdown_text, filename):
        """处理Markdown章节，但不在这里添加文档路径前缀"""
        lines = markdown_text.split('\n')
        sections = []
        current_section = []
        self.heading_stack = []  # 重置标题栈
        current_chapter_path = None
        
        for line in lines:
            level, title = self.parse_heading_level(line)
            
            if level is not None:  # 是标题行
                # 保存之前的section
                if current_section:
                    section_text = '\n'.join(current_section).strip()
                    if section_text:
                        # 不在这里添加文档路径前缀，只保存原始内容和章节路径
                        sections.append((section_text, current_chapter_path))
                    current_section = []
                
                # 更新标题栈
                self.update_heading_stack(level, title)
                
                # 更新当前章节路径
                current_chapter_path = self.get_chapter_path(filename)
                
                # 将标题行添加到新section
                current_section.append(line)
            else:
                # 普通内容行
                current_section.append(line)
        
        # 处理最后一个section
        if current_section:
            section_text = '\n'.join(current_section).strip()
            if section_text:
                # 不在这里添加文档路径前缀，只保存原始内容和章节路径
                sections.append((section_text, current_chapter_path))
        
        return sections
    
    def get_picture_urls(self, sections):
        """获取图片URL"""
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []
        
        from bs4 import BeautifulSoup
        html_content = markdown(text)
        soup = BeautifulSoup(html_content, 'html.parser')
        html_images = [img.get('src') for img in soup.find_all('img') if img.get('src')]
        return html_images
    
    def get_pictures(self, text):
        """下载并打开markdown文本中的所有图片"""
        import requests
        image_urls = self.get_picture_urls(text)
        images = []
        # 查找文本中的所有图片URL
        for url in image_urls:
            try:
                response = requests.get(url, stream=True, timeout=30)
                if response.status_code == 200 and response.headers['Content-Type'].startswith('image/'):
                    img = Image.open(BytesIO(response.content)).convert('RGB')
                    images.append(img)
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
                continue
                    
        return images if images else None

    def mdchapter_merge_with_path(self, sections, chunk_token_num, delimiter, filename):
        """
        使用naive_merge进行分块，然后为每个分块添加文档路径前缀
        """
        from rag.nlp import naive_merge
        
        # 首先使用naive_merge进行分块，传入空的pos确保不添加额外的路径信息
        text_sections = [(text, "") for text, path in sections]  # 确保pos为空字符串
        chunks = naive_merge(text_sections, chunk_token_num, delimiter)
        
        # 为每个分块添加文档路径前缀（只在开头添加一次）
        file_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not file_name:
            file_name = "未命名文档"
            
        prefixed_chunks = []
        for chunk in chunks:
            if chunk.strip():  # 只处理非空分块
                # 清理可能存在的重复文档路径前缀
                cleaned_chunk = self.clean_duplicate_paths(chunk)
                # 只在分块开头添加一次文档路径前缀
                prefixed_chunk = f"【文档路径：{file_name}】\n\n{cleaned_chunk}"
                prefixed_chunks.append(prefixed_chunk)
            else:
                prefixed_chunks.append(chunk)
        
        return prefixed_chunks
    
    def clean_duplicate_paths(self, text):
        """清理文本中重复的文档路径前缀"""
        # 移除所有的【文档路径：...】前缀
        import re
        pattern = r'【文档路径：[^】]*】\s*'
        cleaned_text = re.sub(pattern, '', text)
        return cleaned_text.strip()

    def __call__(self, filename, binary=None):
        """改进的调用方法，使用新的章节处理逻辑"""
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r", encoding='utf-8') as f:
                txt = f.read()
        
        # 提取表格
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        
        # 使用改进的章节处理逻辑
        sections = self.process_markdown_sections(remainder, filename)
        
        # 处理表格
        tbls = []
        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        
        return sections, tbls


# 复制其他类的定义（Docx, Pdf等）保持不变...
class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')
        if not embed:
            return None
        embed = embed[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        except UnicodeDecodeError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph
        
        titles = []
        blocks = []
        
        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"
            
        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith('p'):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(('p', i, p))
                elif block.tag.endswith('tbl'):  # Table
                    blocks.append(('t', i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""
            
        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == 't':
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1
                
        if target_table_pos == -1:
            return ""  # Target table not found
            
        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks)-1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue
                
            if block_type != 'p':
                continue
                
            if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")
        
        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]
            
            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks)-1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue
                        
                    if block_type != 'p':
                        continue
                        
                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:  
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")
                            
                if not found:  # Break if no parent heading is found
                    break
            
            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)
            
        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for i, tb in enumerate(self.doc.tables):
            title = self.__get_nearest_title(i, filename)
            html = "<table>"
            if title:
                html += f"<caption>Table Location: {title}</caption>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            # self._naive_vertical_merge()
            self._concat_downward()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        改进的Markdown章节切片解析器
        支持层次化章节识别和路径标注
        
        特性:
        1. 识别1-6级Markdown标题
        2. 构建章节层次结构
        3. 为每个切片添加文档路径信息
        4. 格式: 【文档路径：文件名 > 1级标题 > 2级标题】
        
        Supported file formats are docx, pdf, excel, txt, markdown.
    """

    is_english = lang.lower() == "english"
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    
    # 获取文件名（不含路径）
    import os
    base_filename = os.path.basename(filename) if filename else "未命名文档"
    
    doc = {
        "docnm_kwd": base_filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", base_filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    section_images = None
    
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        try:
            vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT)
            callback(0.15, "Visual model detected. Attempting to enhance figure extraction...")
        except Exception:
            vision_model = None

        sections, tables = Docx()(filename, binary)

        if vision_model:
            figures_data = vision_figure_parser_figure_data_wrapper(sections)
            try:
                docx_vision_parser = VisionFigureParser(vision_model=vision_model, figures_data=figures_data, **kwargs)
                boosted_figures = docx_vision_parser(callback=callback)
                tables.extend(boosted_figures)
            except Exception as e:
                callback(0.6, f"Visual model error: {e}. Skipping figure parsing enhancement.")

        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
        logging.info("mdchapter_improved_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"
        callback(0.1, "Start to parse.")

        if layout_recognizer == "DeepDOC":
            pdf_parser = Pdf()

            try:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT)
                callback(0.15, "Visual model detected. Attempting to enhance figure extraction...")
            except Exception:
                vision_model = None

            if vision_model:
                sections, tables, figures = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback, separate_tables_figures=True)
                callback(0.5, "Basic parsing complete. Proceeding with figure enhancement...")
                try:
                    pdf_vision_parser = VisionFigureParser(vision_model=vision_model, figures_data=figures, **kwargs)
                    boosted_figures = pdf_vision_parser(callback=callback)
                    tables.extend(boosted_figures)
                except Exception as e:
                    callback(0.6, f"Visual model error: {e}. Skipping figure parsing enhancement.")
                    tables.extend(figures)
            else:
                sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)

            res = tokenize_table(tables, doc, is_english)
            callback(0.8, "Finish parsing.")

        else:
            if layout_recognizer == "Plain Text":
                pdf_parser = PlainParser()
            else:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT, llm_name=layout_recognizer, lang=lang)
                pdf_parser = VisionParser(vision_model=vision_model, **kwargs)

            sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page,
                                          callback=callback)
            res = tokenize_table(tables, doc, is_english)
            callback(0.8, "Finish parsing.")

    elif re.search(r"\.(csv|xlsx?)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = ExcelParser()
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        # 使用改进的Markdown解析器
        markdown_parser = ImprovedMarkdown(int(parser_config.get("chunk_token_num", 128)))
        sections, tables = markdown_parser(base_filename, binary)
        
        # Process images for each section
        section_images = []
        for section_text, chapter_path in sections:
            images = markdown_parser.get_pictures(section_text) if section_text else None
            if images:
                # If multiple images found, combine them using concat_img
                combined_image = reduce(concat_img, images) if len(images) > 1 else images[0]
                section_images.append(combined_image)
            else:
                section_images.append(None)
                
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    if section_images:
        # if all images are None, set section_images to None
        if all(image is None for image in section_images):
            section_images = None

    # 对于markdown文件，使用特殊的合并逻辑来添加文档路径前缀
    if re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        # 创建markdown_parser实例来调用我们的新方法
        if 'markdown_parser' not in locals():
            markdown_parser = ImprovedMarkdown(int(parser_config.get("chunk_token_num", 128)))
        
        if section_images and not all(image is None for image in section_images):
            # 对于有图片的情况，需要先用naive_merge_with_images，然后添加路径前缀
            text_sections = [(text, "") for text, path in sections]  # 转换为标准格式
            chunks, images = naive_merge_with_images(text_sections, section_images,
                                            int(parser_config.get(
                                                "chunk_token_num", 128)), parser_config.get(
                                                "delimiter", "\n!?。；！？"))
            
            # 为每个分块添加文档路径前缀（只在开头添加一次）
            file_name = re.sub(r"\.[a-zA-Z]+$", "", base_filename)
            if not file_name:
                file_name = "未命名文档"
                
            prefixed_chunks = []
            for chunk in chunks:
                if chunk.strip():  # 只处理非空分块
                    # 清理可能存在的重复文档路径前缀
                    cleaned_chunk = markdown_parser.clean_duplicate_paths(chunk)
                    # 只在分块开头添加一次文档路径前缀
                    prefixed_chunk = f"【文档路径：{file_name}】\n\n{cleaned_chunk}"
                    prefixed_chunks.append(prefixed_chunk)
                else:
                    prefixed_chunks.append(chunk)
            
            if kwargs.get("section_only", False):
                return prefixed_chunks
            
            res.extend(tokenize_chunks_with_images(prefixed_chunks, doc, is_english, images))
        else:
            # 没有图片的情况，使用我们的新合并方法
            chunks = markdown_parser.mdchapter_merge_with_path(
                sections, 
                int(parser_config.get("chunk_token_num", 128)), 
                parser_config.get("delimiter", "\n!?。；！？"),
                base_filename
            )
            
            if kwargs.get("section_only", False):
                return chunks

            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    else:
        # 非markdown文件使用原来的逻辑
        if section_images:
            chunks, images = naive_merge_with_images(sections, section_images,
                                            int(parser_config.get(
                                                "chunk_token_num", 128)), parser_config.get(
                                                "delimiter", "\n!?。；！？"))
            if kwargs.get("section_only", False):
                return chunks
            
            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
        else:
            chunks = naive_merge(
                sections, int(parser_config.get(
                    "chunk_token_num", 128)), parser_config.get(
                    "delimiter", "\n!?。；！？"))
            if kwargs.get("section_only", False):
                return chunks

            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    
    logging.info("mdchapter_improved_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)